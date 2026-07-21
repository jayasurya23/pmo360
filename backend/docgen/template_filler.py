"""
Generate meeting minutes by populating the canonical Castillo .docm template
in-place rather than rebuilding from scratch.

The template at templates/_External__Meeting_Minutes_-_Template_-_mm-dd-yy.docm
contains the exact design (Jost font, 14pt red title + date on one row, 18pt
project name with red underline, 14pt red section headings, banded tables, the
template's specific spacing). By cloning the template and rewriting only the
content, we keep every styling detail without trying to mirror it with
python-docx run properties.

The approach:
  1. Open the .docm zip and rewrite its [Content_Types].xml so it loads as a
     .docx (python-docx rejects macro-enabled content type). Drop vbaProject.
  2. Walk the body, locating each section by anchor text from the template's
     sample data (the Nov 7 / Heelstone meeting).
  3. Rewrite paragraph text in place (single-run collapse), or — for the list
     sections and tables — clone the template's first row/paragraph and use
     it as a row template for the new content.
  4. Save as .docx.
"""
from __future__ import annotations

import io
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell

from db.models import Meeting


TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent
    / "templates"
    / "_External__Meeting_Minutes_-_Template_-_mm-dd-yy.docm"
)

# Sample-data anchors in the shipped template — these are what we replace.
ANCHOR_DATE = "November 7, 2025"
ANCHOR_PROJECT = "Snapdragon and Two Blues"
ANCHOR_CLIENT = "Heelstone"
ANCHOR_SCOPE = "Electrical Design, Civil Design and Studies"

SECTION_HEADINGS = (
    "Attendees",
    "Deliverable Timelines",
    "Agenda",
    "Discussion Points",
    "Action Items",
    "Closing Remarks",
)


# ============================================================
# Zip-level conversion: .docm → in-memory .docx stream
# ============================================================
def _docm_as_docx_stream(path: Path) -> io.BytesIO:
    """Repack the .docm so python-docx accepts it. Drops the macro container
    and any .rels Relationship that points at it."""
    import re as _re
    out = io.BytesIO()
    macro_ct = b"application/vnd.ms-word.document.macroEnabled.main+xml"
    docx_ct = b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
    rels_vba_re = _re.compile(
        rb'<Relationship[^>]*vbaProject\.bin[^>]*/>'
    )
    with zipfile.ZipFile(path, "r") as src, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for name in src.namelist():
            if name == "word/vbaProject.bin":
                continue  # drop macros
            data = src.read(name)
            if name == "[Content_Types].xml":
                data = data.replace(macro_ct, docx_ct)
                data = data.replace(
                    b'<Override PartName="/word/vbaProject.bin" ContentType="application/vnd.ms-office.vbaProject"/>',
                    b"",
                )
            elif name.endswith(".rels"):
                # Strip the vbaProject relationship if present
                data = rels_vba_re.sub(b"", data)
            dst.writestr(name, data)
    out.seek(0)
    return out


# ============================================================
# Paragraph helpers
# ============================================================
def _paragraph_text(p) -> str:
    """Visible text of a paragraph, joining all runs."""
    return "".join((r.text or "") for r in p.runs)


def _is_heading(p, heading_text: str) -> bool:
    return _paragraph_text(p).strip() == heading_text


def _find_paragraph_by_text(doc, fragment: str):
    for p in doc.paragraphs:
        if fragment in _paragraph_text(p):
            return p
    return None


def _replace_substring_in_runs(p, old: str, new: str) -> bool:
    """Replace `old` with `new` in paragraph p. Collapses into the first run,
    preserving its formatting. The remaining runs' text is cleared but their
    formatting elements stay — that's fine because empty runs render nothing.
    """
    joined = _paragraph_text(p)
    if old not in joined:
        return False
    new_joined = joined.replace(old, new)
    if not p.runs:
        p.add_run(new_joined)
        return True
    p.runs[0].text = new_joined
    for r in p.runs[1:]:
        r.text = ""
    return True


# ============================================================
# Run-level formatting — rebuild a paragraph's runs so a bold label can sit
# ahead of normal-weight content (matching the PDF). The template otherwise
# collapses everything into one run, which flattens the bold/normal split.
# ============================================================
def _set_run_bold(r, bold: bool) -> None:
    """Force a run's bold state on/off via its rPr, overriding whatever the
    cloned template run carried."""
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = rPr.makeelement(qn(tag), {})
            rPr.append(el)
        el.set(qn("w:val"), "1" if bold else "0")


def _run_set_content(r, text: str) -> None:
    """Replace a run's content with `text`, keeping its rPr. Tab characters
    become <w:tab/> elements so they honour the paragraph's tab stops."""
    for child in list(r):
        if child.tag != qn("w:rPr"):
            r.remove(child)
    parts = (text or "").split("\t")
    for i, part in enumerate(parts):
        if i > 0:
            r.append(r.makeelement(qn("w:tab"), {}))
        if part:
            t = r.makeelement(qn("w:t"), {})
            t.text = part
            t.set(qn("xml:space"), "preserve")
            r.append(t)


def _fill_paragraph_runs(new_p, segments) -> None:
    """Rebuild a paragraph's runs from its first run, one run per
    (text, bold) segment. Preserves the base run's font/size/colour while
    setting bold explicitly, so a bold label can precede normal content."""
    base_runs = new_p.findall(qn("w:r"))
    base = deepcopy(base_runs[0]) if base_runs else None
    new_runs = []
    for text, bold in segments:
        r = deepcopy(base) if base is not None else new_p.makeelement(qn("w:r"), {})
        if bold is not None:
            _set_run_bold(r, bool(bold))
        _run_set_content(r, text)
        new_runs.append(r)
    for r in base_runs:
        new_p.remove(r)
    pPr = new_p.find(qn("w:pPr"))
    if pPr is None:
        for r in new_runs:
            new_p.append(r)
    else:
        anchor = pPr
        for r in new_runs:
            anchor.addnext(r)
            anchor = r


# ============================================================
# Section walking — given a heading, iterate the body elements that follow
# until we hit the next heading.
# ============================================================
def _body_children(doc):
    return list(doc.element.body.iterchildren())


def _section_range(doc, heading_text: str, stop_headings: Iterable[str]):
    """Return (heading_para_element, [next_elements_until_stop])."""
    children = _body_children(doc)
    head_idx = -1
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1]
        if tag != "p":
            continue
        # Match by visible text against the python-docx paragraph wrapper
        for p in doc.paragraphs:
            if p._element is child and _paragraph_text(p).strip() == heading_text:
                head_idx = i
                break
        if head_idx != -1:
            break
    if head_idx == -1:
        return None, []
    stop_set = set(stop_headings)
    section = []
    for child in children[head_idx + 1:]:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            # If this paragraph is itself a known heading, stop
            for p in doc.paragraphs:
                if p._element is child and _paragraph_text(p).strip() in stop_set:
                    return children[head_idx], section
        section.append(child)
    return children[head_idx], section


# ============================================================
# Attendees rewrite — kill existing org lines, clone one as template,
# insert new lines after the heading.
# ============================================================
def _rewrite_attendees(doc, by_org: dict[str, list[str]]):
    head, section = _section_range(doc, "Attendees",
                                   stop_headings=("Deliverable Timelines", "Agenda"))
    if head is None:
        return
    # Existing org-line paragraphs (skip empty / non-paragraph elements)
    existing_paras = []
    for el in section:
        if el.tag.split("}")[-1] == "p":
            for p in doc.paragraphs:
                if p._element is el:
                    if _paragraph_text(p).strip():
                        existing_paras.append(el)
                    break
    if not existing_paras:
        return
    # Keep the first one as a clone template
    template_el = deepcopy(existing_paras[0])
    parent = head.getparent()
    # Remove existing lines (we'll re-insert)
    for el in existing_paras:
        parent.remove(el)
    # Insert new lines right after the heading
    anchor = head
    for org, names in by_org.items():
        new_p = deepcopy(template_el)
        # Bold org label + normal names, matching the PDF (`<b>Org:</b> names`).
        _fill_paragraph_runs(new_p, [(f"{org}: ", True), (names, False)])
        anchor.addnext(new_p)
        anchor = new_p


# ============================================================
# Listing rewrite (Agenda, Discussion Points)
# ============================================================
def _rewrite_listing(doc, heading_text: str, stop_headings: Iterable[str],
                     entries):
    """Rewrite a bulleted listing section (Agenda, Discussion Points).

    `entries` is a list of (indent, segments) where `segments` is a list of
    (text, bold) tuples — each becomes its own run, so a bold label can sit
    ahead of normal-weight content. `indent` (0/1/2) nests sub-bullets.
    """
    head, section = _section_range(doc, heading_text, stop_headings)
    if head is None:
        return
    existing_paras = []
    for el in section:
        if el.tag.split("}")[-1] == "p":
            for p in doc.paragraphs:
                if p._element is el and _paragraph_text(p).strip():
                    existing_paras.append(el)
                    break
    if not existing_paras:
        return
    template_el = deepcopy(existing_paras[0])
    parent = head.getparent()
    for el in existing_paras:
        parent.remove(el)

    from lxml import etree
    anchor = head
    for indent, segments in entries:
        new_p = deepcopy(template_el)

        # Apply paragraph indent (~0.25" per nesting level) so sub-bullets sit
        # below their parent without losing the template's run formatting.
        if indent > 0:
            pPr = new_p.find(qn("w:pPr"))
            if pPr is None:
                pPr = etree.SubElement(new_p, qn("w:pPr"))
                new_p.insert(0, pPr)
            ind = pPr.find(qn("w:ind"))
            if ind is None:
                ind = etree.SubElement(pPr, qn("w:ind"))
            ind.set(qn("w:left"), str(360 * (indent + 1)))

        _fill_paragraph_runs(new_p, segments)
        anchor.addnext(new_p)
        anchor = new_p


# ============================================================
# Table rewrite — keep header row, replace data rows
# ============================================================
def _table_after_heading(doc, heading_text: str):
    """Find the first w:tbl element after the given heading paragraph."""
    head, section = _section_range(doc, heading_text,
                                   stop_headings=SECTION_HEADINGS)
    if head is None:
        return None
    for el in section:
        if el.tag.split("}")[-1] == "tbl":
            for t in doc.tables:
                if t._element is el:
                    return t
    return None


def _remove_section(doc, heading_text: str, stop_headings: Iterable[str]) -> None:
    """Delete a whole section — its heading paragraph and everything up to the
    next known heading. Used to drop optional sections that have no content
    (e.g. Deliverable Timelines when the meeting has no deliverables), matching
    the PDF which simply omits them."""
    head, section = _section_range(doc, heading_text, stop_headings)
    if head is None:
        return
    parent = head.getparent()
    for el in section:
        parent.remove(el)
    parent.remove(head)


def _text_width_twips(doc) -> int:
    """Usable text width (page width minus L/R margins) of the first section,
    in twips (1 inch = 1440 twips = 914400 EMU)."""
    sec = doc.sections[0]
    emu = sec.page_width - sec.left_margin - sec.right_margin
    return int(emu / 914400 * 1440)


def _normalize_table_columns(tbl, max_width_twips: int,
                             col_widths: "list[int] | None" = None) -> None:
    """Force every cell in every row to the table's grid column widths and fix
    the layout, so a cloned/relabelled header cell can't widen a column past
    the page. If the grid is wider than the text area, scale it down to fit.

    `col_widths` overrides the template's grid entirely (e.g. to give the
    Action column more room and shrink the Due column) — it is scaled to fit
    the text width like any other grid."""
    tblEl = tbl._tbl
    grid = tblEl.find(qn("w:tblGrid"))
    if grid is None:
        return
    cols = grid.findall(qn("w:gridCol"))
    if col_widths and len(col_widths) == len(cols):
        widths = [int(w) for w in col_widths]
    else:
        widths = [int(c.get(qn("w:w")) or 0) for c in cols]
    if not widths or sum(widths) <= 0:
        return
    total = sum(widths)
    if total > max_width_twips:
        scale = max_width_twips / total
        widths = [max(1, int(round(w * scale))) for w in widths]
        total = sum(widths)
    # Write the (overridden and/or scaled) widths back into the grid.
    for c, w in zip(cols, widths):
        c.set(qn("w:w"), str(w))
    # Every row's cells snap to the grid column widths.
    for tr in tblEl.findall(qn("w:tr")):
        for i, tc in enumerate(tr.findall(qn("w:tc"))):
            if i >= len(widths):
                break
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = tc.makeelement(qn("w:tcPr"), {})
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = tcPr.makeelement(qn("w:tcW"), {})
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[i]))
            tcW.set(qn("w:type"), "dxa")
    # Fixed layout + explicit table width so Word honours the grid exactly.
    tblPr = tblEl.find(qn("w:tblPr"))
    if tblPr is not None:
        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = tblPr.makeelement(qn("w:tblLayout"), {})
            tblPr.append(layout)
        layout.set(qn("w:type"), "fixed")
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = tblPr.makeelement(qn("w:tblW"), {})
            tblPr.append(tblW)
        tblW.set(qn("w:w"), str(total))
        tblW.set(qn("w:type"), "dxa")


# Status → (cell fill, text color) — mirrors the PDF's status pills
# (docgen/pdf_builder.py STATUS_COLORS / STATUS_TEXT_COLORS) so the Word minutes
# color the Action Items Status column the same way.
_STATUS_FILL = {
    "open":      "1AA6C9",   # blue
    "pending":   "C7BB2E",   # gold
    "completed": "278747",   # green
    "cancelled": "E6E7E8",   # light gray
}
_STATUS_TEXT = "1A1A1A"      # black on all (matches the PDF)


def _shade_tc(tc, fill_hex: str):
    """Set/replace a table cell's background shading (w:shd fill)."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = tc.makeelement(qn("w:tcPr"), {})
        tc.insert(0, tcPr)
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_tc_font_color(tc, hex_color: str):
    """Force the font color of every run in a cell."""
    for r in tc.findall(f".//{qn('w:r')}"):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = r.makeelement(qn("w:rPr"), {})
            r.insert(0, rPr)
        color = rPr.find(qn("w:color"))
        if color is None:
            color = rPr.makeelement(qn("w:color"), {})
            rPr.append(color)
        color.set(qn("w:val"), hex_color)


def _color_action_status(doc, actions):
    """Shade the Action Items Status cell (last column) per status, matching the
    PDF's colored status pills. Data rows follow the action order 1:1."""
    tbl = _table_after_heading(doc, "Action Items")
    if tbl is None:
        return
    for ri, a in enumerate(actions, start=1):
        if ri >= len(tbl.rows):
            break
        fill = _STATUS_FILL.get((getattr(a, "status", "") or "open").lower())
        if not fill:
            continue
        tcs = tbl.rows[ri]._tr.findall(qn("w:tc"))
        if not tcs:
            continue
        _shade_tc(tcs[-1], fill)          # Status is the last column
        _set_tc_font_color(tcs[-1], _STATUS_TEXT)


def _set_cell_text(cell, text: str):
    """Replace cell text while preserving run-level formatting from the first
    run of the first paragraph. Removes extra paragraphs in the cell."""
    if not cell.paragraphs:
        cell.text = str(text)
        return
    p = cell.paragraphs[0]
    runs = p.runs
    if runs:
        runs[0].text = str(text)
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(str(text))
    # Drop any subsequent paragraphs in this cell
    for extra in p._element.getparent().findall(qn("w:p"))[1:]:
        extra.getparent().remove(extra)


def _flatten_cell_sdts(tr) -> None:
    """Unwrap any content control (w:sdt) that wraps table cells in a row,
    replacing it with the plain <w:tc>(s) inside.

    The template's Due / Start / Delivery columns are date-picker content
    controls. Left in place they (a) render their stored placeholder date
    (e.g. 11/10/2025) regardless of the value we write, and (b) hide from
    ``findall('w:tc')`` — so the column count is under-reported and the
    padding logic shifts values into the wrong columns."""
    for sdt in list(tr.findall(qn("w:sdt"))):
        content = sdt.find(qn("w:sdtContent"))
        cells = content.findall(qn("w:tc")) if content is not None else []
        idx = list(tr).index(sdt)
        for j, tc in enumerate(cells):
            tr.insert(idx + j, tc)
        tr.remove(sdt)


def _rewrite_table(doc, heading_text: str, rows_data: list[list[str]],
                   header_labels: "list[str] | None" = None,
                   col_widths: "list[int] | None" = None):
    tbl = _table_after_heading(doc, heading_text)
    if tbl is None or len(tbl.rows) < 2:
        return
    header_el = tbl.rows[0]._tr
    _flatten_cell_sdts(header_el)
    ncols = len(header_el.findall(qn("w:tc")))

    # Optionally (re)label the header. The template's Action Items table ships
    # with a blank Status header; fill it by cloning a styled neighbour cell
    # (so the red shading + white bold carry over) before setting the text.
    if header_labels:
        htcs = header_el.findall(qn("w:tc"))
        for i in range(min(ncols, len(header_labels))):
            tc = htcs[i]
            if not (_Cell(tc, tbl).text or "").strip() and i > 0:
                clone = deepcopy(htcs[i - 1])
                header_el.replace(tc, clone)
                tc = clone
            _set_cell_text(_Cell(tc, tbl), header_labels[i])

    # Clone the first data row as the row template. Some template rows have
    # FEWER <w:tc> than the header (the Action Items rows omit the Due cell, so
    # the Due value would otherwise shift left into the Status cell) — pad the
    # clone up to the header's column count, inserting the missing cell(s)
    # before the last cell with a normal body style.
    template_row_el = deepcopy(tbl.rows[1]._tr)
    _flatten_cell_sdts(template_row_el)
    tcs = template_row_el.findall(qn("w:tc"))
    if tcs and len(tcs) < ncols:
        style_ref = tcs[-2] if len(tcs) >= 2 else tcs[-1]
        while len(tcs) < ncols:
            tcs[-1].addprevious(deepcopy(style_ref))
            tcs = template_row_el.findall(qn("w:tc"))

    # Remove every existing data row (keep header at index 0)
    for row in list(tbl.rows)[1:]:
        tbl._tbl.remove(row._tr)

    # Insert one cloned row per data tuple, writing straight to the raw <w:tc>
    # so value i lands in column i regardless of python-docx grid quirks.
    for row_values in rows_data:
        new_row_el = deepcopy(template_row_el)
        tbl._tbl.append(new_row_el)
        for i, tc in enumerate(new_row_el.findall(qn("w:tc"))):
            _set_cell_text(_Cell(tc, tbl), row_values[i] if i < len(row_values) else "")

    # Snap every cell to the grid so a cloned/relabelled header cell (e.g. the
    # Status column, cloned from the wider Due cell) can't push the table past
    # the page margin and spill the last column outside the table. `col_widths`
    # (when given) also re-proportions the columns.
    _normalize_table_columns(tbl, _text_width_twips(doc), col_widths)


# ============================================================
# Deliverable rows — manual MeetingDeliverable rows, or fallback to the
# upcoming leaf tasks from the project's latest uploaded Schedule.
# ============================================================
def _gather_deliverable_rows(meeting: Meeting) -> list[list[str]]:
    """Build the rows that go into the Deliverable Timelines table from the
    meeting's manually-attached MeetingDeliverable rows. The PM picks these
    on the Review page (optionally choosing items from the project's
    Schedule); we don't auto-pull anything here."""
    rows: list[list[str]] = []
    for idx, md in enumerate(meeting.meeting_deliverables, start=1):
        d = md.deliverable
        rows.append([
            str(idx),
            d.project_segment or meeting.project.name,
            d.task,
            d.start_status or "In Progress",
            d.delivery_date.strftime("%m/%d/%Y") if d.delivery_date else "",
        ])
    return rows


# ============================================================
# Public entry point
# ============================================================
def generate_meeting_minutes_from_template(meeting: Meeting) -> bytes:
    """Open the canonical .docm template, populate with this meeting's data,
    return clean .docx bytes."""
    stream = _docm_as_docx_stream(TEMPLATE_PATH)
    doc = Document(stream)

    # --- Date in the title row ---
    new_date = meeting.meeting_date.strftime("%B %d, %Y")
    title_p = _find_paragraph_by_text(doc, ANCHOR_DATE)
    if title_p:
        _replace_substring_in_runs(title_p, ANCHOR_DATE, new_date)

    # --- Project name ---
    proj_name = meeting.project.name
    for p in doc.paragraphs:
        if _paragraph_text(p).strip() == ANCHOR_PROJECT:
            _replace_substring_in_runs(p, ANCHOR_PROJECT, proj_name)
            break

    # --- Client value (bold label + normal value, matching the PDF) ---
    if meeting.project.client:
        for p in doc.paragraphs:
            if _paragraph_text(p).strip().startswith("Client:"):
                _fill_paragraph_runs(
                    p._element,
                    [("Client:\t", True), (meeting.project.client.name, False)],
                )
                break

    # --- Scope value: bold label + normal value when set; otherwise drop the
    #     row entirely (the PDF only prints Scope when the project has one). ---
    for p in doc.paragraphs:
        if _paragraph_text(p).strip().startswith("Scope:"):
            if meeting.project.scope:
                _fill_paragraph_runs(
                    p._element,
                    [("Scope:\t", True), (meeting.project.scope, False)],
                )
            else:
                p._element.getparent().remove(p._element)
            break

    # --- Attendees grouped by org ---
    by_org: dict[str, list] = {}
    for a in meeting.attendees:
        by_org.setdefault(a.organization or "Other", []).append(a)
    org_to_names = {
        org: ", ".join(f"{a.full_name} ({a.initials})" for a in attendees)
        for org, attendees in by_org.items()
    }
    # Always call the rewriters so the template's sample data is replaced
    # (or cleared, if we don't have any new data to substitute in).
    _rewrite_attendees(doc, org_to_names)

    # --- Deliverable Timelines table (omit the whole section when there are no
    #     deliverables, matching the PDF which renders it only when populated) ---
    deliv_rows = _gather_deliverable_rows(meeting)
    if deliv_rows:
        _rewrite_table(doc, "Deliverable Timelines", deliv_rows)
    else:
        _remove_section(doc, "Deliverable Timelines", SECTION_HEADINGS)

    # --- Agenda (plain bullets, normal weight — matches the PDF) ---
    agenda_entries = [
        (0, [(a.text or "", False)])
        for a in sorted(meeting.agenda_items, key=lambda x: x.order_index)
    ]
    _rewrite_listing(doc, "Agenda", ("Discussion Points",), agenda_entries)

    # --- Discussion points: bold label + normal content, flattened to
    #     (indent, segments) so nested sub-points indent (matches the PDF). ---
    def _dp_segments(dp):
        label = (getattr(dp, "label", "") or "").strip()
        content = dp.content or ""
        if label:
            return [(f"{label}: ", True), (content, False)]
        return [(content, False)]

    disc_entries: list = []

    def _walk(dp, level: int):
        disc_entries.append((level, _dp_segments(dp)))
        for sub in sorted(getattr(dp, "sub_points", []) or [], key=lambda s: s.order_index):
            _walk(sub, level + 1)

    # Only iterate top-level points (parent_id IS NULL); sub_points come via _walk
    top_level = [dp for dp in meeting.discussion_points if getattr(dp, "parent_id", None) is None]
    for dp in sorted(top_level, key=lambda d: d.order_index):
        _walk(dp, 0)
    _rewrite_listing(doc, "Discussion Points", ("Action Items",), disc_entries)

    # --- Action Items table ---
    act_rows = []
    for idx, a in enumerate(meeting.raised_actions, start=1):
        act_rows.append([
            str(idx),
            a.text,
            a.owner or "",
            a.due_date.strftime("%m/%d/%Y") if a.due_date else "",
            a.status.capitalize(),
        ])
    # Column widths (twips, ~7.75" text area): wide Action, roomier Owner so
    # names like "Roashaael Mary John" don't over-wrap, and a compact Due.
    #   #=460  Action=6100  Owner=1900  Due=1200  Status=1500  (sum 11160)
    _rewrite_table(doc, "Action Items", act_rows,
                   header_labels=["#", "Action", "Owner", "Due", "Status"],
                   col_widths=[460, 6100, 1900, 1200, 1500])
    _color_action_status(doc, meeting.raised_actions)

    # --- Closing Remarks ---
    # Keep template's "Thank you to everyone..." text as-is, unless the meeting
    # has explicit closing remarks set.
    if meeting.closing_remarks:
        for p in doc.paragraphs:
            t = _paragraph_text(p).strip()
            if t.startswith("Thank you to everyone"):
                _replace_substring_in_runs(p, t, meeting.closing_remarks)
                break

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
