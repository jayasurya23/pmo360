"""
Generate Castillo-branded Word .docx for meeting minutes.

Matches the heading style and layout of the template Gary uses:
- Red bold title "Meeting Minutes"
- Red underlined section headings (Attendees, Deliverable Timelines, Agenda,
  Discussion Points, Action Items, Closing Remarks)
- Red header bar on tables, banded rows
- Jost-style font (Helvetica fallback)
"""
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from db.models import Meeting, Schedule
from config import BrandColors


# ============================================================
# Logo asset paths
# ============================================================
_LOGO_DIR = Path(__file__).resolve().parent.parent / "assets" / "logo"
_PMO360_LOGO = _LOGO_DIR / "pmo360_logo.png"
_CASTILLO_LOGO_COLOR = _LOGO_DIR / "Castillo_logo_color.png"


def _add_pmo360_logo(doc, width_inches: float = 2.0) -> None:
    """Insert the PMO 360 logo as a centered inline image at the document's
    current position. Falls back to no-op if the file is missing — callers
    can still render the text title alongside."""
    if not _PMO360_LOGO.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(str(_PMO360_LOGO), width=Inches(width_inches))


def _add_castillo_logo_inline(p, width_inches: float = 1.2) -> None:
    """Append a small inline Castillo Engineering color logo to an existing
    paragraph (used in the document header / corner). No-op when missing."""
    if not _CASTILLO_LOGO_COLOR.exists():
        return
    run = p.add_run()
    run.add_picture(str(_CASTILLO_LOGO_COLOR), width=Inches(width_inches))


# Fixed agenda topics used in every Castillo pre-meeting coordination doc.
# The third tuple entry is the 30-minute time allocation; `_scale_topics`
# doubles every numeric "Nm" entry when the meeting is 60-minute. "Open" and
# any non-numeric times pass through verbatim.
PREMEETING_AGENDA_TOPICS = [
    ("Welcome and Introductions",          "Project Manager", "2m"),
    ("Review of Previous Action Items",    "Project Manager", "5m"),
    ("Schedule Update and Change Log",     "Project Manager", "5m"),
    ("Status of Active Deliverables",      "Project Manager", "5m"),
    ("Risks and Constraints",              "Project Manager", "5m"),
    ("Required Decisions",                 "Project Manager", "5m"),
    ("Open Discussion",                    "All",             "Open"),
    ("Summary of Actions and Next Steps",  "Project Manager", "5m"),
]


def _scale_topics(topics, duration_minutes: int):
    factor = max(1, int(duration_minutes or 30)) / 30
    if factor == 1:
        return topics
    out = []
    for topic, presenter, time_str in topics:
        new = time_str
        s = (time_str or "").strip()
        if s.lower().endswith("m") and s[:-1].isdigit():
            new = f"{int(s[:-1]) * int(factor)}m"
        out.append((topic, presenter, new))
    return out

# Discipline buckets used in the Discussion Points and Previous Week Recap sections
AGENDA_DISCUSSION_DISCIPLINES = ["Civil", "Electrical", "Structural", "General"]


# RGBColor versions of brand hexes
def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


RED = _rgb(BrandColors.RED)
NEAR_BLACK = _rgb(BrandColors.NEAR_BLACK)
DARK_GRAY = _rgb(BrandColors.DARK_GRAY)
WHITE = _rgb("#ffffff")
# Castillo template body color — slightly warmer than the UI's near-black.
# Pulled from the actual .docm template's run properties.
TEMPLATE_BODY = _rgb("#333132")
BODY_FONT = "Jost"


def _shade_cell(cell, hex_color: str):
    """Apply a fill color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _add_heading(doc, text: str):
    """Castillo template section heading: 14pt red, Jost, NOT bold,
    with a thin red bottom border."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = False
    run.font.size = Pt(14)
    run.font.color.rgb = RED
    run.font.name = BODY_FONT
    # Bottom border for the heading
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 0.75pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BrandColors.RED.lstrip("#"))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def _body_run(paragraph, text: str, bold: bool = True, size_pt: int = 10):
    """Body run matching the template's default (10pt bold #333132 Jost)."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = TEMPLATE_BODY
    run.font.name = BODY_FONT
    return run


def generate_meeting_minutes_docx(meeting: Meeting, output_path: Optional[Path] = None) -> bytes:
    """Returns the docx bytes; also writes to output_path if given."""
    doc = Document()

    # ---- Page setup ----
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ---- Title row ("Meeting Minutes" + date on same line, both 14pt bold red) ----
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Meeting Minutes")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RED
    title_run.font.name = BODY_FONT
    # right-aligned date on the same paragraph via a right tab stop
    p_pr = title_p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")     # 6.5 inches in twentieths of a point
    tabs.append(tab)
    p_pr.append(tabs)
    title_p.add_run("\t")
    date_run = title_p.add_run(meeting.meeting_date.strftime("%B %d, %Y"))
    date_run.bold = True
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RED
    date_run.font.name = BODY_FONT

    # ---- Project block ----
    proj_p = doc.add_paragraph()
    proj_run = proj_p.add_run(meeting.project.name)
    proj_run.bold = True
    proj_run.font.size = Pt(18)
    proj_run.font.color.rgb = TEMPLATE_BODY
    proj_run.font.name = BODY_FONT
    # Red bottom border
    p_pr = proj_p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BrandColors.RED.lstrip("#"))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    if meeting.project.client:
        kv = doc.add_paragraph()
        _body_run(kv, "Client:  ")
        _body_run(kv, meeting.project.client.name)
    if meeting.project.scope:
        kv = doc.add_paragraph()
        _body_run(kv, "Scope:  ")
        _body_run(kv, meeting.project.scope)

    # ---- Attendees ----
    _add_heading(doc, "Attendees")
    by_org = {}
    for a in meeting.attendees:
        by_org.setdefault(a.organization or "Other", []).append(a)
    for org, attendees in by_org.items():
        line = doc.add_paragraph()
        _body_run(line, f"{org}: ")
        names = ", ".join(f"{a.full_name} ({a.initials})" for a in attendees)
        _body_run(line, names)

    # ---- Deliverable Timelines ----
    if meeting.meeting_deliverables:
        _add_heading(doc, "Deliverable Timelines")
        tbl = doc.add_table(rows=1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tbl.rows[0].cells
        for i, label in enumerate(["#", "Project", "Task", "Start", "Delivery"]):
            cell = hdr[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
            _shade_cell(cell, BrandColors.RED)

        for idx, md in enumerate(meeting.meeting_deliverables, start=1):
            d = md.deliverable
            row = tbl.add_row().cells
            cells_text = [
                str(idx),
                d.project_segment or meeting.project.name,
                d.task,
                d.start_status or "In Progress",
                d.delivery_date.strftime("%m/%d/%Y") if d.delivery_date else "",
            ]
            for c, t in zip(row, cells_text):
                c.text = ""
                _body_run(c.paragraphs[0], t)

    # ---- Agenda ----
    if meeting.agenda_items:
        _add_heading(doc, "Agenda")
        for item in sorted(meeting.agenda_items, key=lambda a: a.order_index):
            p = doc.add_paragraph(style="List Bullet")
            _body_run(p, item.text)

    # ---- Discussion Points ----
    if meeting.discussion_points:
        _add_heading(doc, "Discussion Points")
        for dp in sorted(meeting.discussion_points, key=lambda d: d.order_index):
            p = doc.add_paragraph(style="List Bullet")
            if dp.label:
                _body_run(p, f"{dp.label}: ")
            _body_run(p, dp.content)

    # ---- Action Items ----
    if meeting.client_facing_actions:
        _add_heading(doc, "Action Items")
        tbl = doc.add_table(rows=1, cols=5)
        hdr = tbl.rows[0].cells
        for i, label in enumerate(["#", "Action", "Owner", "Due", "Status"]):
            cell = hdr[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
            _shade_cell(cell, BrandColors.RED)

        for idx, a in enumerate(meeting.client_facing_actions, start=1):
            row = tbl.add_row().cells
            cells_text = [
                str(idx),
                a.text,
                a.owner or "",
                a.due_date.strftime("%m/%d/%Y") if a.due_date else "",
                a.status.capitalize(),
            ]
            for c, t in zip(row, cells_text):
                c.text = ""
                _body_run(c.paragraphs[0], t)

    # ---- Closing Remarks ----
    _add_heading(doc, "Closing Remarks")
    closing = meeting.closing_remarks or (
        "Thank you to everyone for attending this meeting. "
        "Your time is very much appreciated."
    )
    p = doc.add_paragraph()
    _body_run(p, closing)

    # ---- Footer line ----
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = footer_p.add_run("Castillo Engineering  |  Project Management Office  |  Confidential")
    f.font.size = Pt(9)
    f.font.color.rgb = DARK_GRAY

    # ---- Save ----
    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_bytes(data)
    return data


# ============================================================
# Pre-meeting coordination agenda
# ============================================================
def _make_table(doc, headers: list[str], col_widths_in: Optional[list[float]] = None):
    """Create a table with a red-shaded header row matching the Castillo template."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, label in enumerate(headers):
        cell = hdr[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = BODY_FONT
        _shade_cell(cell, BrandColors.RED)
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    return tbl


def _add_subheading(doc, text: str):
    """Smaller heading used for discipline buckets within Discussion Points."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    run.font.name = BODY_FONT
    return p


def generate_premeeting_agenda_docx(
    meeting: Meeting,
    prior_meeting: Optional[Meeting] = None,
    carry_actions: Optional[list] = None,
    carry_deliverables: Optional[list] = None,
    schedule: Optional[Schedule] = None,
    output_path: Optional[Path] = None,
    # ---- New: editable sections built up on the Next Agenda UI ----
    disciplines: Optional[list] = None,
    dp_by_discipline: Optional[dict] = None,
    recap_by_discipline: Optional[dict] = None,
    risks: Optional[list] = None,
    decisions: Optional[list] = None,
    schedule_changes: Optional[list] = None,
    meeting_duration_minutes: int = 30,
    schedule_version_override: Optional[str] = None,
) -> bytes:
    """Generate the Castillo pre-meeting coordination agenda .docx.

    `meeting` is the upcoming draft Meeting (date + project + attendees).
    `prior_meeting` is the most recently completed meeting — used to populate
    the Previous Week Recap discussion buckets when `recap_by_discipline` isn't
    supplied.
    `carry_actions` and `carry_deliverables` are passed pre-fetched so this
    function stays free of repository imports.
    `schedule` (optional) provides the "Current Schedule Version" line.

    The Next Agenda UI assembles the editable sections and passes them in:
      - `disciplines` — ordered list of section names for Discussion Points
        and Previous Week Recap. Defaults to AGENDA_DISCUSSION_DISCIPLINES.
      - `dp_by_discipline` — dict[discipline -> list[ParsedDiscussionPoint]]
        with nested `sub_points`. Rendered under Discussion Points.
      - `recap_by_discipline` — same shape, rendered under Previous Week Recap.
        If None and `prior_meeting` is given, falls back to extracting from
        the prior meeting's discussion_points.
      - `risks` — list[dict] with keys description / impact / likelihood /
        mitigation / owner.
      - `decisions` — list[dict] with keys decision / description /
        impact_if_not / required_by (date | None) / owner.
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ---- Title (PMO 360 logo + "Pre Meeting Coordination" + date) ----
    # The logo replaces the leading text-rendered PMO 360 branding. The
    # subtitle stays so it's clear which deliverable this is.
    _add_pmo360_logo(doc, width_inches=2.0)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Pre Meeting Coordination")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RED
    title_run.font.name = BODY_FONT

    date_p = doc.add_paragraph()
    date_run = date_p.add_run(meeting.meeting_date.strftime("%B %d, %Y"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = DARK_GRAY

    # ---- Project block ----
    proj_p = doc.add_paragraph()
    proj_run = proj_p.add_run(meeting.project.name)
    proj_run.bold = True
    proj_run.font.size = Pt(15)
    proj_run.font.color.rgb = NEAR_BLACK
    p_pr = proj_p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BrandColors.RED.lstrip("#"))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    if meeting.project.client:
        kv = doc.add_paragraph()
        b = kv.add_run("Client:  ")
        b.bold = True
        kv.add_run(meeting.project.client.name)
    if meeting.project.scope:
        kv = doc.add_paragraph()
        b = kv.add_run("Scope:  ")
        b.bold = True
        kv.add_run(meeting.project.scope)

    # ---- Attendees (carried forward from prior meeting) ----
    _add_heading(doc, "Attendees")
    by_org = {}
    for a in meeting.attendees:
        by_org.setdefault(a.organization or "Other", []).append(a)
    if by_org:
        for org, attendees in by_org.items():
            line = doc.add_paragraph()
            b = line.add_run(f"{org}: ")
            b.bold = True
            names = ", ".join(a.full_name for a in attendees)
            line.add_run(names)
    else:
        doc.add_paragraph("(attendee list will be confirmed before the meeting)")

    # ---- Agenda (fixed 8-row table) ----
    _add_heading(doc, "Agenda")
    dur_p = doc.add_paragraph()
    dur_lbl = dur_p.add_run("Meeting duration: ")
    dur_lbl.bold = True
    dur_lbl.font.size = Pt(10)
    dur_val = dur_p.add_run(f"{int(meeting_duration_minutes or 30)} min")
    dur_val.font.size = Pt(10)
    tbl = _make_table(doc, ["#", "Topic", "Presenter", "Time Allocation"],
                      col_widths_in=[0.4, 4.0, 1.5, 1.2])
    scaled_topics = _scale_topics(PREMEETING_AGENDA_TOPICS, meeting_duration_minutes)
    for idx, (topic, presenter, time) in enumerate(scaled_topics, start=1):
        row = tbl.add_row().cells
        row[0].text = str(idx)
        row[1].text = topic
        row[2].text = presenter
        row[3].text = time

    # ---- Deliverable Timelines (carry-forward) ----
    _add_heading(doc, "Deliverable Timelines")
    version_line = doc.add_paragraph()
    _override = (schedule_version_override or "").strip()
    _version_label = (
        _override if _override
        else (schedule.version if schedule
              else (meeting.project.schedule_version or "—"))
    )
    v_run = version_line.add_run(
        f"Current Schedule Version: {_version_label}"
    )
    v_run.bold = True
    v_run.font.size = Pt(10)

    tbl = _make_table(doc, ["#", "Project", "Task", "Start", "Delivery"],
                      col_widths_in=[0.4, 1.6, 3.2, 1.0, 1.0])
    deliverables = carry_deliverables or []
    if deliverables:
        for idx, d in enumerate(deliverables, start=1):
            row = tbl.add_row().cells
            row[0].text = str(idx)
            row[1].text = d.project_segment or meeting.project.name
            row[2].text = d.task
            row[3].text = d.start_status or "In Progress"
            row[4].text = d.delivery_date.strftime("%m/%d/%Y") if d.delivery_date else ""
    else:
        # 4 empty rows for the PM to fill in (matches the template's empty state)
        for idx in range(1, 5):
            row = tbl.add_row().cells
            row[0].text = str(idx)
            row[3].text = "In Progress"

    # ---- Discussion Points (manual, per-discipline) ----
    active_disciplines = disciplines or AGENDA_DISCUSSION_DISCIPLINES
    _add_heading(doc, "Discussion Points")
    dp_map = dp_by_discipline or {}

    def _write_dp_tree(items, level: int = 0) -> None:
        """Render a tree of ParsedDiscussionPoint (or DiscussionPoint) items
        as a `List Bullet` paragraph; sub-points sit one indent deeper."""
        if not items:
            doc.add_paragraph(style="List Bullet").add_run("")
            return
        for dp in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            label = getattr(dp, "label", "") or ""
            content = getattr(dp, "content", "") or ""
            if label:
                lbl = p.add_run(f"{label}: ")
                lbl.bold = True
            p.add_run(content)
            kids = getattr(dp, "sub_points", None) or []
            if kids:
                _write_dp_tree(kids, level + 1)

    for discipline in active_disciplines:
        _add_subheading(doc, discipline)
        _write_dp_tree(dp_map.get(discipline) or [])

    # ---- Previous Week Recap ----
    _add_heading(doc, "Previous Week Recap")
    if recap_by_discipline is None:
        # Backwards-compat: bucket the prior meeting's discussion points by
        # their stored discipline (mapping anything unrecognized to General).
        recap_by_discipline = {d: [] for d in active_disciplines}
        if prior_meeting:
            for dp in sorted(prior_meeting.discussion_points,
                             key=lambda d: d.order_index):
                disc = (dp.discipline or "General").capitalize()
                if disc not in recap_by_discipline:
                    disc = "General" if "General" in recap_by_discipline else active_disciplines[-1]
                recap_by_discipline[disc].append(dp)
    for discipline in active_disciplines:
        _add_subheading(doc, discipline)
        _write_dp_tree(recap_by_discipline.get(discipline) or [])

    # ---- Schedule Change Log ----
    _add_heading(doc, "Schedule Change Log")
    tbl = _make_table(doc, [
        "Project", "Task / Milestone", "Previous Date", "Updated Date",
        "Change Description", "Reason for Change", "Impact on Overall Schedule",
    ])
    sc_rows = [r for r in (schedule_changes or [])
               if any((r.get(k) or "").strip()
                      for k in ("project", "task", "previous_date",
                                "updated_date", "change_description",
                                "reason_for_change", "impact"))]
    if sc_rows:
        for r in sc_rows:
            row = tbl.add_row().cells
            row[0].text = (r.get("project") or "").strip()
            row[1].text = (r.get("task") or "").strip()
            row[2].text = (r.get("previous_date") or "").strip()
            row[3].text = (r.get("updated_date") or "").strip()
            row[4].text = (r.get("change_description") or "").strip()
            row[5].text = (r.get("reason_for_change") or "").strip()
            row[6].text = (r.get("impact") or "").strip()
    else:
        for _ in range(5):
            tbl.add_row()

    # ---- Risks and Constraints ----
    _add_heading(doc, "Risks and Constraints")
    tbl = _make_table(doc, [
        "Description", "Impact", "Likelihood", "Mitigation Strategy", "Owner",
    ])
    risk_rows = [r for r in (risks or [])
                 if any((r.get(k) or "").strip()
                        for k in ("description", "impact", "likelihood",
                                  "mitigation", "owner"))]
    if risk_rows:
        for r in risk_rows:
            row = tbl.add_row().cells
            row[0].text = (r.get("description") or "").strip()
            row[1].text = (r.get("impact") or "").strip()
            row[2].text = (r.get("likelihood") or "").strip()
            row[3].text = (r.get("mitigation") or "").strip()
            row[4].text = (r.get("owner") or "").strip()
    else:
        for _ in range(5):
            tbl.add_row()

    # ---- Required Decisions ----
    _add_heading(doc, "Required Decisions")
    tbl = _make_table(doc, [
        "Decision", "Description", "Impact if Not Decided",
        "Required By Date", "Decision Owner",
    ])
    decision_rows = []
    for d in (decisions or []):
        has_text = any((d.get(k) or "").strip()
                       for k in ("decision", "description",
                                 "impact_if_not", "owner"))
        has_date = bool(d.get("required_by"))
        if has_text or has_date:
            decision_rows.append(d)
    if decision_rows:
        for d in decision_rows:
            row = tbl.add_row().cells
            row[0].text = (d.get("decision") or "").strip()
            row[1].text = (d.get("description") or "").strip()
            row[2].text = (d.get("impact_if_not") or "").strip()
            rb = d.get("required_by")
            if rb:
                try:
                    row[3].text = rb.strftime("%m/%d/%Y")
                except AttributeError:
                    row[3].text = str(rb)
            row[4].text = (d.get("owner") or "").strip()
    else:
        for _ in range(5):
            tbl.add_row()

    # ---- Open action items carried into this meeting ----
    if carry_actions:
        _add_heading(doc, "Open Action Items")
        tbl = _make_table(doc, ["#", "Action", "Owner", "Due", "Status"],
                          col_widths_in=[0.4, 4.0, 0.9, 1.0, 0.9])
        for idx, a in enumerate(carry_actions, start=1):
            row = tbl.add_row().cells
            row[0].text = str(idx)
            row[1].text = a.text
            row[2].text = a.owner or ""
            row[3].text = a.due_date.strftime("%m/%d/%Y") if a.due_date else ""
            row[4].text = a.status.capitalize()

    # ---- Closing Remarks ----
    _add_heading(doc, "Closing Remarks")
    doc.add_paragraph(
        "Thank you to everyone for attending this meeting. "
        "Your time is very much appreciated."
    )

    # ---- Footer ----
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = footer_p.add_run("Castillo Engineering  |  Project Management Office  |  Confidential")
    f.font.size = Pt(9)
    f.font.color.rgb = DARK_GRAY

    # ---- Save ----
    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_bytes(data)
    return data
